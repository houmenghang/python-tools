from docx import Document
from docx.shared import Inches
import json

class build_doc():
    def __init__(self):
        self.document = Document()

    def add_title(self, doc_name= 'chengyu_story'):
        self.document.add_heading(doc_name, 0)
    
    def add_content(self, title, text):
        self.document.add_heading(title, level=1)
        # self.document.add_picture(image, width=Inches(5))
        self.document.add_paragraph(text)
        
    def add_page(self):
        self.document.add_page_break()

    def save_doc(self, name= 'chengyu_story.docx'):
        self.document.save(name)

data = json.load(open('chengyu_data\chengyu_content.json', 'r', encoding='utf-8'))

doc = build_doc()
i = 0
for data_info in data:
    i += 1
    id = str(i).rjust(4, '0')
    doc.add_title(id)
    doc.add_content(data_info['title'], data_info['content'].replace('\r\n\t\t\t\t　　', '\n'))

doc.save_doc()